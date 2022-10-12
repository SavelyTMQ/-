import math

from main import main_sup, clean_dataset, affinity
import docx
import numpy as np
import os
import pandas as pd
from playsound import *
import datetime
import time
import matplotlib.pyplot as plt
from openpyxl import load_workbook

import seaborn as sns
from sklearn.decomposition import PCA
import random
import scipy.stats as sps

### нормализуем данные [0, 1] и сохраняем их в папке Данные по кластерам
def data_normalization():
    directory = '/Users/User/Desktop/ДИПЛОМ/Данные по кластерам2'
    files = os.listdir(directory)

    Q = []
    T = pd.DataFrame()

    for i in range(len(files)):
        data = pd.read_csv('Данные по кластерам2/' + files[i])
        T = pd.concat([T, data])
        Q.append(data.shape[0])

    F = T.to_numpy()

    Max = []
    Min = np.array([np.min(F[:, i]) for i in range(len(F[0]))])
    F = F + abs(Min)

    Max = [np.max(F[:, i]) for i in range(len(F[0]))]

    F = [F[:, i] / Max[i] for i in range(len(Max))]

    F = np.transpose(F)

    print(Max)

    print(Min)

    """u = 0
    for i in range(len(files)):
        print(i + 1)
        # ['BENIGN', 'DDoS', 'PortScan', 'Bot', 'Infiltration', 'Web Attack � Brute Force',
        # 'Web Attack � XSS', 'Web Attack � Sql Injection', 'FTP-Patator', 'SSH-Patator',
        # 'DoS slowloris', 'DoS Slowhttptest', 'DoS Hulk', 'DoS GoldenEye', 'Heartbleed']

        W = pd.DataFrame(F[u:u + Q[i]])

        W.to_csv('Данные по кластерам/' + files[i], index=False)
        u = u + Q[i]"""

    return

### находим все виды атак
def attack_types():
    directory = '/Users/User/Desktop/ДИПЛОМ/MachineLearningCVE'
    files = os.listdir(directory)

    C = []
    for j in files:

        data = pd.read_csv('MachineLearningCVE/' + j)
        data = data.to_numpy()
        for i in data:
            if not i[78] in C:
                C.append(i[78])

    return C

### на основе корреляции удаляем лишние столбцы
def separation_of_attacks():
    directory = '/Users/User/Desktop/ДИПЛОМ/MachineLearningCVE'
    files = os.listdir(directory)
    data = pd.read_csv('MachineLearningCVE/' + files[0])

    BENIGN = pd.DataFrame()
    DDoS = pd.DataFrame()
    PortScan = pd.DataFrame()
    Bot = pd.DataFrame()
    Infiltration = pd.DataFrame()
    Brute_Force = pd.DataFrame()
    XSS = pd.DataFrame()
    Sql_Injection = pd.DataFrame()
    FTP_Patator = pd.DataFrame()
    SSH_Patator = pd.DataFrame()
    DoS_slowloris = pd.DataFrame()
    DoS_Slowhttptest = pd.DataFrame()
    DoS_Hulk = pd.DataFrame()
    DoS_GoldenEye = pd.DataFrame()
    Heartbleed = pd.DataFrame()

    T = np.array([1, 1, 1, 0, 1, 0, 1, 1, 0, 0, 1, 1, 0, 0, 1, 1, 0, 0, 0, 1, 0, 0, 0, 0, 1,
                  0, 0, 0, 0, 0, 1, 0, 1, 0, 1, 1, 0, 1, 0, 0, 0, 0, 0, 1, 0, 1, 1, 0, 0, 0,
                  0, 1, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 1, 0, 0, 1, 1, 0, 0, 0,
                  1, 0, 0])

    for i in range(len(files)):
        data = pd.read_csv('MachineLearningCVE/' + files[i])

        BENIGN = pd.concat([BENIGN, data[data[' Label'].str.contains("BENIGN")]])
        DDoS = pd.concat([DDoS, data[data[' Label'].str.contains("DDoS")]])
        PortScan = pd.concat([PortScan, data[data[' Label'].str.contains("PortScan")]])
        Bot = pd.concat([Bot, data[data[' Label'].str.contains("Bot")]])
        Infiltration = pd.concat([Infiltration, data[data[' Label'].str.contains("Infiltration")]])
        Brute_Force = pd.concat([Brute_Force, data[data[' Label'].str.contains("Web Attack � Brute Force")]])
        XSS = pd.concat([XSS, data[data[' Label'].str.contains("Web Attack � XSS")]])
        Sql_Injection = pd.concat([Sql_Injection, data[data[' Label'].str.contains("Web Attack � Sql Injection")]])
        FTP_Patator = pd.concat([FTP_Patator, data[data[' Label'].str.contains("FTP-Patator")]])
        SSH_Patator = pd.concat([SSH_Patator, data[data[' Label'].str.contains("SSH-Patator")]])
        DoS_slowloris = pd.concat([DoS_slowloris, data[data[' Label'].str.contains("DoS slowloris")]])
        DoS_Slowhttptest = pd.concat([DoS_Slowhttptest, data[data[' Label'].str.contains("DoS Slowhttptest")]])
        DoS_Hulk = pd.concat([DoS_Hulk, data[data[' Label'].str.contains("DoS Hulk")]])
        DoS_GoldenEye = pd.concat([DoS_GoldenEye, data[data[' Label'].str.contains("DoS GoldenEye")]])
        Heartbleed = pd.concat([Heartbleed, data[data[' Label'].str.contains("Heartbleed")]])

    del BENIGN[' Label']
    del DDoS[' Label']
    del PortScan[' Label']
    del Bot[' Label']
    del Infiltration[' Label']
    del Brute_Force[' Label']
    del XSS[' Label']
    del Sql_Injection[' Label']
    del FTP_Patator[' Label']
    del SSH_Patator[' Label']
    del DoS_slowloris[' Label']
    del DoS_Slowhttptest[' Label']
    del DoS_Hulk[' Label']
    del DoS_GoldenEye[' Label']
    del Heartbleed[' Label']

    BENIGN = clean_dataset(BENIGN)
    DDoS = clean_dataset(DDoS)
    PortScan = clean_dataset(PortScan)
    Bot = clean_dataset(Bot)
    Infiltration = clean_dataset(Infiltration)
    Brute_Force = clean_dataset(Brute_Force)
    XSS = clean_dataset(XSS)
    Sql_Injection = clean_dataset(Sql_Injection)
    FTP_Patator = clean_dataset(FTP_Patator)
    SSH_Patator = clean_dataset(SSH_Patator)
    DoS_slowloris = clean_dataset(DoS_slowloris)
    DoS_Slowhttptest = clean_dataset(DoS_Slowhttptest)
    DoS_Hulk = clean_dataset(DoS_Hulk)
    DoS_GoldenEye = clean_dataset(DoS_GoldenEye)
    Heartbleed = clean_dataset(Heartbleed)

    BENIGN.iloc[0:int(len(BENIGN) / 3), T == 1].to_csv('./Данные по кластерам2/BENIGN1.csv', index=False)
    BENIGN.iloc[int(len(BENIGN) / 3):2 * int(len(BENIGN) / 3), T == 1].to_csv('./Данные по кластерам2/BENIGN2.csv', index=False)
    BENIGN.iloc[2 * int(len(BENIGN) / 3):, T == 1].to_csv('./Данные по кластерам2/BENIGN3.csv', index=False)
    DDoS.iloc[:, T == 1].to_csv('./Данные по кластерам2/DDoS.csv', index=False)
    PortScan.iloc[:, T == 1].to_csv('./Данные по кластерам2/PortScan.csv', index=False)
    Bot.iloc[:, T == 1].to_csv('./Данные по кластерам2/Bot.csv', index=False)
    Infiltration.iloc[:, T == 1].to_csv('./Данные по кластерам2/Infiltration.csv', index=False)
    Brute_Force.iloc[:, T == 1].to_csv('./Данные по кластерам2/Brute_Force.csv', index=False)
    XSS.iloc[:, T == 1].to_csv('./Данные по кластерам2/XSS.csv', index=False)
    Sql_Injection.iloc[:, T == 1].to_csv('./Данные по кластерам2/Sql_Injection.csv', index=False)
    FTP_Patator.iloc[:, T == 1].to_csv('./Данные по кластерам2/FTP_Patator.csv', index=False)
    SSH_Patator.iloc[:, T == 1].to_csv('./Данные по кластерам2/SSH_Patator.csv', index=False)
    DoS_slowloris.iloc[:, T == 1].to_csv('./Данные по кластерам2/DoS_slowloris.csv', index=False)
    DoS_Slowhttptest.iloc[:, T == 1].to_csv('./Данные по кластерам2/DoS_Slowhttptest.csv', index=False)
    DoS_Hulk.iloc[:, T == 1].to_csv('./Данные по кластерам2/DoS_Hulk.csv', index=False)
    DoS_GoldenEye.iloc[:, T == 1].to_csv('./Данные по кластерам2/DoS_GoldenEye.csv', index=False)
    Heartbleed.iloc[:, T == 1].to_csv('./Данные по кластерам2/Heartbleed.csv', index=False)

###
def main():
    #name = 'MachineLearningCVE/Friday-WorkingHours-Afternoon-DDos.pcap_ISCX.csv'
    for u in range(8, 17):
        attack = pd.DataFrame()
        no_attack = pd.DataFrame()

        directory = '/Users/User/Desktop/ДИПЛОМ/Данные по кластерам'
        files = os.listdir(directory)


        for i in range(len(files)):
            data = pd.read_csv('Данные по кластерам/' + files[i])
            #print(files[i][:-4])

            if i == u:
                no_attack = pd.concat([no_attack, data])
                print(files[i])
            else:
                attack = pd.concat([attack, data])

        np_attack = attack.to_numpy()
        np_no_attack = no_attack.to_numpy()

        #print(np.shape(np_no_attack))
        #return
        # кол-во антител для начальной популяции
        N_benign = 200
        N_Bot = 300
        N_Brute_Force = 100
        N_DDOS = 400
        N_DoS_GoldenEye = 400
        N_DoS_Hulk = 500
        N_DoS_Slowhttptest = 500
        N_DoS_slowloris = 300
        N_FTP_Patator = 500
        N_Heartbleed = 4
        N_Infiltration = 23
        N_PortScan = 200
        N_Sql_Injection = 11
        N_SSH_Patator = 200
        N_XSS = 400



        N = [N_benign, N_Bot, N_Brute_Force, N_DDOS, N_DoS_GoldenEye, N_DoS_Hulk, N_DoS_Slowhttptest,
             N_DoS_slowloris, N_FTP_Patator, N_Heartbleed, N_Infiltration, N_PortScan, N_Sql_Injection,
             N_SSH_Patator, N_XSS]

        # ct - пороговое значение аффиности для клонирования
        ct_benign = 0.075
        ct_Bot = 0.5

        ct = [ct_benign, ct_Bot]

        # nt - пороговое значение аффиности для добавления в популяцию
        nt_benign = 1
        nt_Bot = 1

        nt = [nt_benign, nt_Bot]

        B_benign = 0.08
        B_Bot = 0.04
        B_Brute_Force = 0.07
        B_DDOS = 0.08
        B_DoS_GoldenEye = 0.05
        B_DoS_Hulk = 0.09
        B_DoS_Slowhttptest = 0.08
        B_DoS_slowloris = 0.03
        B_FTP_Patator = 0.1
        B_Heartbleed = 0.07
        B_Infiltration = 0.07
        B_PortScan = 0.08
        B_Sql_Injection = 0.03
        B_SSH_Patator = 0.08
        B_XSS = 0.1

        B = [B_benign, B_Bot, B_Brute_Force, B_DDOS, B_DoS_GoldenEye, B_DoS_Hulk, B_DoS_Slowhttptest,
             B_DoS_slowloris, B_FTP_Patator, B_Heartbleed, B_Infiltration, B_PortScan, B_Sql_Injection,
             B_SSH_Patator, B_XSS]

        # границы мутации
        mut_benign = [-0.01, 0.01]
        mut_Bot = [-0.00005, 0.00005]
        mut_Brute_Force = [-0.00001, 0.00001]
        mut_DDOS = [-0.00001, 0.00001]
        mut_DoS_GoldenEye = [-0.01, 0.01]
        mut_DoS_Hulk = [-0.03, 0.03]
        mut_DoS_Slowhttptest = [-0.0001, 0.0001]
        mut_DoS_slowloris = [-0.001, 0.001]
        mut_FTP_Patator = [-0.00005, 0.00005]
        mut_Heartbleed = [-0.005, 0.005]
        mut_Infiltration = [-0.00001, 0.00001]
        mut_PortScan = [-0.0005, 0.0005]
        mut_Sql_Injection = [-0.01, 0.01]
        mut_SSH_Patator = [-0.01, 0.01]
        mut_XSS = [-0.0001, 0.0001]

        mut = [mut_benign, mut_Bot, mut_Brute_Force, mut_DDOS, mut_DoS_GoldenEye, mut_DoS_Hulk, mut_DoS_Slowhttptest,
               mut_DoS_slowloris, mut_FTP_Patator, mut_Heartbleed, mut_Infiltration, mut_PortScan, mut_Sql_Injection,
               mut_SSH_Patator, mut_XSS]

        # тестовая выборка
        M = 1000

        # пороговая аффиность
        aff_benign = 0.7
        aff_Bot = 0.125
        aff_Brute_Force = 0.2
        aff_DDOS = 0.25
        aff_DoS_GoldenEye = 0.25
        aff_DoS_Hulk = 0.4
        aff_DoS_Slowhttptest = 0.4
        aff_DoS_slowloris = 0.7
        aff_FTP_Patator = 0.45
        aff_Heartbleed = 0.85
        aff_Infiltration = 0.8
        aff_PortScan = 0.8
        aff_Sql_Injection = 0.55
        aff_SSH_Patator = 0.55
        aff_XSS = 0.1

        aff = [aff_benign, aff_Bot, aff_Brute_Force, aff_DDOS, aff_DoS_GoldenEye, aff_DoS_Hulk, aff_DoS_Slowhttptest,
             aff_DoS_slowloris, aff_FTP_Patator, aff_Heartbleed, aff_Infiltration, aff_PortScan, aff_Sql_Injection,
             aff_SSH_Patator, aff_XSS]

        #N = [10, 25, 50, 75, 100, 200, 300, 400, 500]
        #N = [3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23, 24, 25, 26, 27, 28,
        #     29, 30, 31, 32, 33, 34]

        #N = [4, 10, 11, 16]

        #A = [0.1, 0.2, 0.25, 0.3, 0.4, 0.45, 0.5, 0.55, 0.6, 0.65, 0.7, 0.75, 0.8, 0.85]
        #B = [0.1, 0.09, 0.08, 0.07, 0.06, 0.05, 0.04, 0.03]

        #mut = [[-0.03, 0.03], [-0.01, 0.01], [-0.005, 0.005], [-0.001, 0.001], [-0.0005, 0.0005],
        #       [-0.0001, 0.0001], [-0.00005, 0.00005], [-0.00001, 0.00001]]

        ct = [1.0, 0.75, 0.5, 0.25]
        nt = [1.0, 0.75, 0.5, 0.25]

        for t in range(len(ct)):


            for h in range(len(nt)):
                print(f"Сейчас используется {ct[t]}, цикл {t} и {nt[h]}, цикл {h}")



                P = np.array([])
                Obch = np.array([])
                Test = np.array([])
                TSC = np.array([])

                # запускаем алгоритм 10 раз
                for i in range(10):
                    print(i)
                    p, len_C, c_mid, te, ob, ts = main_sup(np_attack, np_no_attack, N[u - 2], ct[t], nt[h], 10, B[u - 2], mut[u - 2], M, aff[u - 2])

                    P = np.append(P, p)
                    Obch = np.append(Obch, ob)
                    Test = np.append(Test, te)
                    TSC = np.append(TSC, ts)
                    #C = np.vstack((C, c_mid))
                    #P = np.append(P, p)

                time = datetime.datetime.now()

                doc = docx.Document(f'Тесты/Подбор параметров атак/{files[u][:-4]}/5.ct и nt.docx')

                # добавляем первый параграф
                doc.add_paragraph(f'Дата {time}')

                doc.add_paragraph('N = {} ct = {} nt = {} TSC = {} B = {} mut = {} M = {} '
                                  'aff = {} популяция = {}'
                                  .format(N[u - 2], ct[t], nt[h], TSC, B[u - 2], mut[u - 2], M, aff[u - 2], len_C))

                table = doc.add_table(rows=4, cols=10)
                # применяем стиль для таблицы
                table.style = 'Table Grid'

                # заполняем таблицу данными
                for col in range(10):
                    # получаем ячейку таблицы
                    cell = table.cell(0, col)
                    # записываем в ячейку данные
                    cell.text = str(P[col])

                    # получаем ячейку таблицы
                    cell = table.cell(1, col)
                    # записываем в ячейку данные
                    cell.text = str(Obch[col])

                    # получаем ячейку таблицы
                    cell = table.cell(2, col)
                    # записываем в ячейку данные
                    cell.text = str(Test[col])

                    # получаем ячейку таблицы
                    cell = table.cell(3, col)
                    # записываем в ячейку данные
                    cell.text = str(TSC[col])


                doc.add_paragraph('p = {}'.format(np.sum(P) / len(P)))

                doc.save(f'Тесты/Подбор параметров атак/{files[u][:-4]}/5.ct и nt.docx')

                #playsound('Filatov_Karas_GAYAZOV_BROTHER_-_Poshla_zhara_72992182.mp3')

    playsound('Filatov_Karas_GAYAZOV_BROTHER_-_Poshla_zhara_72992182.mp3')
    return
    #print("Вероятность угадывания ", P)
    #return
    #doc = docx.Document('/Tests/Only_good_data/Test1.docx')
    #doc = docx.Document('Тесты/Только хорошие данные/Test1.docx')
    doc = docx.Document('ДЛЯ ЭС.docx')

    # добавляем первый параграф
    doc.add_paragraph('N = {} ct = {} nt = {} TSC = {} B = {} mut = {} M = {} aff = {} популяция = {}'
                      .format(N, ct, nt, TSC, B, mut, M, aff, len_C))


    table = doc.add_table(rows=1, cols=10)
    # применяем стиль для таблицы
    table.style = 'Table Grid'

    # заполняем таблицу данными
    for row in range(1):
        for col in range(10):
            # получаем ячейку таблицы
            cell = table.cell(row, col)
            # записываем в ячейку данные
            cell.text = str(P[col])

    table = doc.add_table(rows=10, cols=25)
    print(np.shape(C))

    # заполняем таблицу данными
    for row in range(10):
        for col in range(25):
            # получаем ячейку таблицы
            cell = table.cell(row, col)
            # записываем в ячейку данные
            cell.text = str(C[1 + row, col])

    doc.add_paragraph('p = {}'.format(np.sum(P) / len(P)))

    doc.save('ДЛЯ ЭС.docx')
    #doc.save('Тесты/Только хорошие данные/Test1.docx')

### корреляция данных, ищем пары столбцов с высокой корреляцией и удаляем один из них
def sample_cleaning():
    directory = '/Users/User/Desktop/ДИПЛОМ/MachineLearningCVE'
    files = os.listdir(directory)
    T = pd.DataFrame()

    for i in range(len(files)):
        print(i + 1)
        data = pd.read_csv('MachineLearningCVE/' + files[i])

        T = pd.concat([T, data])
    del T[' Label']

    T = clean_dataset(T)
    Q = []
    M = T.to_numpy()

    for i in range(len(M[0])):
        if np.sum(M[:, i]) == 0:
            Q.append(i)

    M = np.delete(M, Q, axis=1)
    Mat = np.transpose(M)
    C = np.corrcoef(Mat)

    f = open('корреляция.txt', 'w')
    f.write(str(C))
    f.close()

    K = np.ones((len(C)), dtype=int)
    i = 0
    while i < (len(C) - 1):
        j = i + 1
        while j < len(C) and K[i] != 0:
            if abs(C[i, j]) >= 0.5:
                K[j] = 0
            j += 1
        i += 1

    K = list(K)
    for i in range(len(Q)):
        K.insert(Q[i], 0)
    return K


###
def analysis():
    directory = '/Users/User/Desktop/ДИПЛОМ/MachineLearningCVE'
    files = os.listdir(directory)

    T = np.array([1, 1, 1, 0, 1, 0, 1, 1, 0, 0, 1, 1, 0, 0, 1, 1, 0, 0, 0, 1, 0, 0, 0, 0, 1,
                  0, 0, 0, 0, 0, 1, 0, 1, 0, 1, 1, 0, 1, 0, 0, 0, 0, 0, 1, 0, 1, 1, 0, 0, 0,
                  0, 1, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 1, 0, 0, 1, 1, 0, 0, 0,
                  1, 0, 0, 0])

    Q = np.zeros((78, 511807))

    data = pd.read_csv('MachineLearningCVE/' + files[0])
    data = np.array(list(data))
    print(data[T == 1])

###
def D():
    NO = np.zeros((1, 25))
    NO = np.vstack((NO, [1.36754837e-02,  8.19024038e-03,  1.37205472e-05,  4.03214164e-06,
                     2.15037164e-03,  1.49749340e-02,  5.18447632e-03,  2.45460167e-02,
                     1.12387444e-01,  3.37077926e-01,  1.32926700e-04,  7.77016194e-04,
                     4.87973454e-03,  4.12827243e-06,  9.99843712e-01,  9.94586197e-01,
                     4.49856960e-03,  1.44676398e-03, - 9.31991687e-07,  2.27456981e-02,
                     6.26069228e-03,  3.21325695e-03,  7.73170314e-05,  3.48793148e-05,
                     3.27190466e-04]))

    Bot = [1.27193619e-01,  8.42751747e-03,  2.22778404e-05,  4.59223092e-06,
     3.85320402e-04,  2.30926827e-05,  3.63708219e-04,  2.05364724e-03,
     1.11923552e-01,  3.33533409e-01,  5.41483750e-06,  4.19153654e-03,
     - 7.27961915e-06, - 1.94600582e-06,  9.99858133e-01,  9.94593436e-01,
     3.01166812e-04,  4.78847718e-06, - 4.32434584e-07,  9.90550124e-01,
     6.35644626e-03,  3.71392565e-04, - 2.85766580e-06, - 2.97681685e-06,
     1.09797346e-06]
    NO = np.vstack((NO, Bot))

    Brute_Force = [1.22038643e-03,  4.87587401e-02,  1.85996854e-05,  1.48406259e-06,
     1.91309721e-05, - 8.73229846e-08,  4.42898516e-05, - 6.75616417e-08,
     1.11921310e-01,  3.33339333e-01,  9.40453330e-07,  7.57972955e-06,
     9.05232657e-09, - 1.02545858e-07,  9.99855778e-01,  9.94591927e-01,
     8.88539368e-06,  4.39521865e-08,  5.11220627e-08,  9.99095336e-01,
     8.26505655e-06,  4.41145325e-01, - 1.98283727e-08, - 1.53792952e-07,
     - 2.30086051e-07]
    NO = np.vstack((NO, Brute_Force))

    DDoS = [ 1.22069423e-0,  1.71450274e-02,  1.66929052e-05,  2.60883936e-06,
    8.07091041e-04,  6.81684134e-06,  5.82791234e-01,  5.40361473e-07,
    1.12102153e-01,  3.33367957e-01,  1.91024958e-06,  2.09178801e-06,
    -2.16960593e-06,  1.20586674e-06,  9.99856939e-01,  9.94591781e-01,
    5.44200487e-05,  9.73111317e-06,  2.18995885e-06,  9.69072430e-01,
    6.22033098e-03, 3.49762149e-03,  9.28076098e-04, 1.52033615e-06,
    5.15322036e-04]
    NO = np.vstack((NO, DDoS))

    DoS_GoldenEye = [ 1.22084553e-03,  5.02698243e-02,  4.06496542e-05,  3.81887190e-05,
  1.97309451e-02, -2.87399645e-07,  2.98864918e-01,  1.56226592e-07,
  1.11921818e-01,  3.33333496e-01,  1.48340315e-04,  1.95138018e-04,
 -1.05718776e-07, -8.40403089e-08,  9.99855866e-01,  9.94592280e-01,
  5.08970192e-07, -1.95846226e-07, -1.75559879e-07, 9.99542976e-01,
  9.70695589e-07,  3.59928588e-03,  5.07767161e-06,  3.10030361e-07,
  2.05906870e-06]
    NO = np.vstack((NO, DoS_GoldenEye))

    DoS_Hulk = [1.22092871e-03,  7.74950692e-01,  2.87889043e-05,  2.53359348e-05,
  1.29018803e-02,  3.64838266e-06,  3.33882011e-01, -8.70267723e-08,
  1.11922122e-01,  3.38742745e-01,  4.10203760e-08,  5.00721223e-07,
  1.53088654e-07, -1.31788974e-08,  9.99855889e-01,  9.94591841e-01,
  7.37606558e-06,  5.51871058e-01,  4.00107054e-07,  3.10099137e-03,
  5.98983620e-03,  3.44155188e-03,  6.01244054e-05,  2.17223856e-07,
  5.84252269e-04]
    NO = np.vstack((NO, DoS_Hulk))

    DoS_Slowhttptest = [ 1.22069090e-03,  5.90522748e-01,  3.41434034e-05,  6.82548325e-06,
  3.49675611e-03,  3.71666302e-05,  5.41286946e-06,  6.72002890e-08,
  1.11927517e-01,  3.33335776e-01,  6.94633010e-03,  6.95095405e-03,
  1.23762984e-04, -3.11338681e-07,  9.99855793e-01,  9.94592052e-01,
  3.80168661e-06, -6.51357830e-08,  2.48523345e-07,  9.99803526e-01,
  1.40873178e-06, 6.28935606e-04,  5.30817232e-02,  2.83052524e-06,
  1.32167825e-01]
    NO = np.vstack((NO, DoS_Slowhttptest))

    DoS_slowloris = [1.22091580e-03,  8.53241438e-01,  6.93029515e-05,  1.95598704e-04,
     9.27135987e-03,  1.14350534e-05,  3.09972672e-04, - 2.08653909e-07,
     1.11923405e-01,  3.33338726e-01,  1.72751329e-05,  1.17281400e-03,
     1.98879198e-03,  4.54431524e-07,  9.99856849e-01,  9.94593333e-01,
     8.43705072e-06, - 1.19819465e-06,  5.58878371e-07,  9.97913727e-01,
     1.20979213e-05,  1.30600045e-04,  2.90232285e-02,  6.08250916e-02,
     2.43380293e-01]
    NO = np.vstack((NO, DoS_slowloris))

    FTP_Patator = [3.18033070e-04,  3.62589737e-03,  1.65399990e-05, - 7.07205336e-07,
     5.86505406e-04,  1.77586025e-06,  8.99188814e-05, - 5.51666856e-07,
     1.11945094e-01,  3.35315544e-01,  1.63385568e-06,  4.27628590e-06,
     9.49251978e-01, - 7.36180652e-07,  9.99855962e-01,  9.94589625e-01,
     1.83049672e-03,  2.62308306e-06, - 1.54877278e-06,  5.07303393e-02,
     3.25295815e-04,  1.95847524e-04, - 1.90484827e-06,  4.83824714e-07,
     1.53373521e-06]
    NO = np.vstack((NO, FTP_Patator))

    Heartbleed = [6.77500572e-03, 9.22330920e-01, 1.17615912e-02, 9.69746300e-04,
     1.82217420e-01, 0.00000000e+00, 7.61644091e-01, 0.00000000e+00,
     1.11949358e-01, 3.33340093e-01, 1.13636350e-07, 9.99999900e-08,
     0.00000000e+00, 0.00000000e+00, 9.99858390e-01, 9.94648168e-01,
     8.58853210e-06, 0.00000000e+00, 0.00000000e+00, 9.09090909e-02,
     0.00000000e+00, 4.23362038e-03, 1.83223140e-06, 0.00000000e+00,
     0.00000000e+00]
    NO = np.vstack((NO, Heartbleed))

    Infiltration = [6.77500572e-03, 6.53397719e-01, 3.78240909e-03, 2.92035487e-02,
     4.07053004e-02, 2.15053763e-03, 2.13773682e-03, 1.61141805e-03,
     1.11929755e-01, 3.33803344e-01, 6.85286957e-05, 1.40844754e-02,
     5.55555556e-01, 0.00000000e+00, 9.99856339e-01, 9.94607303e-01,
     3.72415690e-06, 0.00000000e+00, 0.00000000e+00, 1.66666667e-01,
     4.09544160e-03, 2.64434814e-02, 3.68454998e-02, 2.29115837e-02,
     1.19172350e-01]
    NO = np.vstack((NO, Infiltration))

    PortScan = [3.65872564e-02,  3.58411585e-05,  9.95197168e-06, - 1.58928463e-06,
     2.17076436e-06,  2.38599972e-05,  3.17048792e-04,  2.07036395e-03,
     1.11965574e-01,  3.38905105e-01,  1.17806634e-06,  3.06695604e-05,
     6.03979550e-07, - 1.24356432e-06,  9.99854555e-01,  9.94592912e-01,
     8.35672948e-03, - 1.79845461e-06,  6.57094822e-07,  9.99968759e-01,
     6.40810814e-03,  1.91498652e-04, - 8.55130173e-07,  4.36925554e-07,
     8.03185633e-07]
    NO = np.vstack((NO, PortScan))

    Sql_Injection = [1.22072175e-03, 2.39200919e-02, 1.84183612e-05, 2.44961240e-05,
     1.27316680e-02, 0.00000000e+00, 5.88008680e-02, 0.00000000e+00,
     1.11921234e-01, 3.35465684e-01, 4.57142804e-07, 6.19841208e-07,
     0.00000000e+00, 0.00000000e+00, 9.99855827e-01, 9.94592016e-01,
     3.17507436e-03, 0.00000000e+00, 0.00000000e+00, 5.71428571e-01,
     4.88400488e-03, 3.60834031e-03, 0.00000000e+00, 0.00000000e+00,
     0.00000000e+00]
    NO = np.vstack((NO, Sql_Injection))


    SSH_Patator = [3.35861152e-04,  7.30638960e-02,  7.24972376e-05,  1.07768892e-04,
     1.79190043e-02, - 3.27798948e-07,  3.47277166e-02,  2.83339334e-07,
     1.11921041e-01,  3.36609059e-01,  2.13449310e-07, 1.85292841e-06,
     2.92368530e-05, - 3.05825785e-08,  9.99855802e-01,  9.94592531e-01,
     1.87930959e-03, - 2.07784821e-07,  2.33946654e-07,  6.94945827e-01,
     6.28017235e-03,  3.70773377e-03,  1.11531920e-07,  5.27439347e-08,
     - 2.21519180e-07]
    NO = np.vstack((NO, SSH_Patator))

    XSS = [1.22048381e-03,  4.72880619e-02,  1.84756114e-05,  8.05652241e-08,
     1.25939708e-07, - 7.56665872e-08,  3.26936848e-06, - 4.04349030e-08,
     1.11920755e-01,  3.33333657e-01,  7.35054024e-07,  7.11614342e-06,
     2.08854044e-07,  1.56518873e-07,  9.99855673e-01,  9.94591988e-01,
     3.11190337e-08, - 1.05097021e-08, - 1.55318038e-07,  9.99965853e-01,
     - 1.32340020e-07,  4.41885593e-01, - 1.77696653e-08,  8.47318874e-08,
     - 1.34327464e-07]
    NO = np.vstack((NO, XSS))


    #data = pd.read_csv('ЭС.csv').to_numpy()

    #BENIGN.iloc[0:int(len(BENIGN) / 3), T == 1].to_csv('./Данные по кластерам2/BENIGN1.csv', index=False)

    #data = np.vstack((data, DDoS))
    data = pd.DataFrame(NO)
    data.to_csv('ЭС.csv', index=False)

    #BENIGN = pd.concat([BENIGN, data[data[' Label'].str.contains("BENIGN")]])

###
def check():
    attack = pd.DataFrame()

    directory = '/Users/User/Desktop/ДИПЛОМ/Данные по кластерам'
    files = os.listdir(directory)

    NO = np.zeros((1, 25))

    Bot = np.zeros((1, 25))

    Brute_Force = np.zeros((1, 25))

    DDoS = np.zeros((1, 25))

    DoS_GoldenEye = np.zeros((1, 25))

    DoS_Hulk = np.zeros((1, 25))

    DoS_Slowhttptest = np.zeros((1, 25))

    DoS_slowloris = np.zeros((1, 25))

    FTP_Patator = np.zeros((1, 25))

    Heartbleed = np.zeros((1, 25))

    Infiltration = np.zeros((1, 25))

    PortScan = np.zeros((1, 25))

    Sql_Injection = np.zeros((1, 25))

    SSH_Patator = np.zeros((1, 25))

    XSS = np.zeros((1, 25))

    Q = [NO, Bot, Brute_Force, DDoS, DoS_GoldenEye, DoS_Hulk, DoS_Slowhttptest,
         DoS_slowloris, FTP_Patator, Heartbleed, Infiltration, PortScan,
         Sql_Injection, SSH_Patator, XSS]

    for i in range(len(files)):
        data = pd.read_csv('Данные по кластерам/' + files[i])

        attack = data.to_numpy()

        ind = np.arange(len(attack))
        np.random.shuffle(ind)
        attack = attack[ind]

        if i < 3:
            Q[0] = np.vstack((Q[0], attack[:int(0.1 * len(attack))]))
        else:
            Q[i - 2] = np.vstack((Q[i - 2], attack[:int(0.1 * len(attack))]))
            #print(Q[i - 2])

    data = pd.read_csv('ЭС.csv')
    T = data.to_numpy()[1:]

    p1 = 0
    p2 = 0

    Sum = 0
    Sum = sum([len(i) for i in Q]) - len(Q)


    for i in range(len(Q)):
        for j in range(len(Q[i])):
            C = np.array([affinity(z, Q[i][j]) for z in T])

            c = np.argmin(C)
            # первый случай, хотим определить атака или нет
            if c == i == 0:
                p1 += 1
            if c != 0 and i != 0:
                p1 += 1

            if c == i:
                p2 += 1

    print(p1 / Sum)

    print(p2 / Sum)

###
def check_sample():

    #attack = pd.DataFrame()
    #no_attack = pd.DataFrame()

    NO = pd.DataFrame()
    Bot = pd.DataFrame()
    Brute_Force = pd.DataFrame()
    DDoS = pd.DataFrame()
    DoS_GoldenEye = pd.DataFrame()
    DoS_Hulk = pd.DataFrame()
    DoS_Slowhttptest = pd.DataFrame()
    DoS_slowloris = pd.DataFrame()
    FTP_Patator = pd.DataFrame()
    Heartbleed = pd.DataFrame()
    Infiltration = pd.DataFrame()
    PortScan = pd.DataFrame()
    Sql_Injection = pd.DataFrame()
    SSH_Patator = pd.DataFrame()
    XSS = pd.DataFrame()

    directory = '/Users/User/Desktop/ДИПЛОМ/Данные по кластерам'
    files = os.listdir(directory)

    for i in range(len(files)):
        data = pd.read_csv('Данные по кластерам/' + files[i])

        if i < 3:
            NO = pd.concat([NO, data])

        if i == 3:
            Bot = pd.concat([Bot, data])

        if i == 4:
            Brute_Force = pd.concat([Brute_Force, data])

        if i == 5:
            DDoS = pd.concat([DDoS, data])

        if i == 6:
            DoS_GoldenEye = pd.concat([DoS_GoldenEye, data])

        if i == 7:
            DoS_Hulk = pd.concat([DoS_Hulk, data])

        if i == 8:
            DoS_Slowhttptest = pd.concat([DoS_Slowhttptest, data])

        if i == 9:
            DoS_slowloris = pd.concat([DoS_slowloris, data])

        if i == 10:
            FTP_Patator = pd.concat([FTP_Patator, data])

        if i == 11:
            Heartbleed = pd.concat([Heartbleed, data])

        if i == 12:
            Infiltration = pd.concat([Infiltration, data])

        if i == 13:
            PortScan = pd.concat([PortScan, data])

        if i == 14:
            Sql_Injection = pd.concat([Sql_Injection, data])

        if i == 15:
            SSH_Patator = pd.concat([SSH_Patator, data])

        if i == 16:
            XSS = pd.concat([XSS, data])

    NO = NO.to_numpy()
    Bot = Bot.to_numpy()
    Brute_Force = Brute_Force.to_numpy()
    DDoS = DDoS.to_numpy()
    DoS_GoldenEye = DoS_GoldenEye.to_numpy()
    DoS_Hulk = DoS_Hulk.to_numpy()
    DoS_Slowhttptest = DoS_Slowhttptest.to_numpy()
    DoS_slowloris = DoS_slowloris.to_numpy()
    FTP_Patator = FTP_Patator.to_numpy()
    Heartbleed = Heartbleed.to_numpy()
    Infiltration = Infiltration.to_numpy()
    PortScan = PortScan.to_numpy()
    Sql_Injection = Sql_Injection.to_numpy()
    SSH_Patator = SSH_Patator.to_numpy()
    XSS = XSS.to_numpy()

    NO = np.mean(NO, axis=0)
    Bot = np.mean(Bot, axis=0)
    Brute_Force = np.mean(Brute_Force, axis=0)
    DDoS = np.mean(DDoS, axis=0)
    DoS_GoldenEye = np.mean(DoS_GoldenEye, axis=0)
    DoS_Hulk = np.mean(DoS_Hulk, axis=0)
    DoS_Slowhttptest = np.mean(DoS_Slowhttptest, axis=0)
    DoS_slowloris = np.mean(DoS_slowloris, axis=0)
    FTP_Patator = np.mean(FTP_Patator, axis=0)
    Heartbleed = np.mean(Heartbleed, axis=0)
    Infiltration = np.mean(Infiltration, axis=0)
    PortScan = np.mean(PortScan, axis=0)
    Sql_Injection = np.mean(Sql_Injection, axis=0)
    SSH_Patator = np.mean(SSH_Patator, axis=0)
    XSS = np.mean(XSS, axis=0)

    Attack = [NO, Bot, Brute_Force, DDoS, DoS_GoldenEye, DoS_Hulk, DoS_Slowhttptest, DoS_slowloris,
              FTP_Patator, Heartbleed, Infiltration, PortScan, Sql_Injection, SSH_Patator, XSS]

    F = np.ones((len(Attack), len(Attack)))

    for i in range(len(Attack)):
        for j in range(len(Attack)):
            F[i][j] = affinity(Attack[i], Attack[j])

    doc = docx.Document('Тесты/Тестирование выборки/Связь разных атак.docx')

    # добавляем первый параграф
    doc.add_paragraph(f'Дата {time}')

    table = doc.add_table(rows=15, cols=15)
    # применяем стиль для таблицы
    table.style = 'Table Grid'

    # заполняем таблицу данными
    for row in range(15):
        for col in range(15):
            # получаем ячейку таблицы
            cell = table.cell(row, col)
            # записываем в ячейку данные
            cell.text = str(F[row, col])

    doc.save('Тесты/Тестирование выборки/Связь разных атак.docx')

### выборочное среднее и выборочная дисперсия
def D_E():
    directory = '/Users/User/Desktop/ДИПЛОМ/Данные по кластерам'
    files = os.listdir(directory)

    NO = np.zeros((1, 25))
    Bot = np.zeros((1, 25))
    Brute_Force = np.zeros((1, 25))
    DDoS = np.zeros((1, 25))
    DoS_GoldenEye = np.zeros((1, 25))
    DoS_Hulk = np.zeros((1, 25))
    DoS_Slowhttptest = np.zeros((1, 25))
    DoS_slowloris = np.zeros((1, 25))
    FTP_Patator = np.zeros((1, 25))
    Heartbleed = np.zeros((1, 25))
    Infiltration = np.zeros((1, 25))
    PortScan = np.zeros((1, 25))
    Sql_Injection = np.zeros((1, 25))
    SSH_Patator = np.zeros((1, 25))
    XSS = np.zeros((1, 25))

    Q = [NO, Bot, Brute_Force, DDoS, DoS_GoldenEye, DoS_Hulk, DoS_Slowhttptest,
         DoS_slowloris, FTP_Patator, Heartbleed, Infiltration, PortScan,
         Sql_Injection, SSH_Patator, XSS]

    no_attack = np.zeros((1, 25))

    for i in range(len(files)):
        data = pd.read_csv('Данные по кластерам/' + files[i])

        attack = data.to_numpy()

        if i < 3:
            no_attack = np.vstack((no_attack, attack))

        if i == 2:
            print(len(no_attack))
            no_attack = np.delete(no_attack, 0, axis=0)
            print(len(no_attack))
            Q[0] = np.mean(no_attack, axis=0)

        elif i > 2:
            print(len(attack))
            Q[i - 2] = np.mean(attack, axis=0)

    print(np.shape(Q))

###
def model_training():
    #name = 'MachineLearningCVE/Friday-WorkingHours-Afternoon-DDos.pcap_ISCX.csv'
    T = [0.7, 0.67, 0.9, 0.6, 0.82, 0.7, 0.85, 0.72, 0.75, 0.942, 0.845, 0.9, 0.7, 0.66, 0.964]
    u = 10
    while u < 17:
        print(u)
        attack = pd.DataFrame()
        no_attack = pd.DataFrame()

        directory = '/Users/User/Desktop/ДИПЛОМ/Данные по кластерам'
        files = os.listdir(directory)

        print(files[u])
        for i in range(len(files)):
            data = pd.read_csv('Данные по кластерам/' + files[i])
            #print(files[i][:-4])

            if i < 3 and u < 3:
                no_attack = pd.concat([no_attack, data])

            elif i == u:
                no_attack = pd.concat([no_attack, data])

            elif i != u:
                attack = pd.concat([attack, data])

        np_attack = attack.to_numpy()
        np_no_attack = no_attack.to_numpy()
        print(len(np_no_attack))

        # кол-во антител для начальной популяции
        N_benign = 200
        N_Bot = 300
        N_Brute_Force = 100
        N_DDOS = 400
        N_DoS_GoldenEye = 400
        N_DoS_Hulk = 500
        N_DoS_Slowhttptest = 500
        N_DoS_slowloris = 300
        N_FTP_Patator = 500
        N_Heartbleed = 4
        N_Infiltration = 23
        N_PortScan = 200
        N_Sql_Injection = 11
        N_SSH_Patator = 200
        N_XSS = 400

        N = [N_benign, N_Bot, N_Brute_Force, N_DDOS, N_DoS_GoldenEye, N_DoS_Hulk, N_DoS_Slowhttptest,
             N_DoS_slowloris, N_FTP_Patator, N_Heartbleed, N_Infiltration, N_PortScan, N_Sql_Injection,
             N_SSH_Patator, N_XSS]

        # ct - пороговое значение аффиности для клонирования
        ct_benign = 0.075
        ct_Bot = 0.5
        ct_Brute_Force = 0.1
        ct_DDOS = 0.75
        ct_DoS_GoldenEye = 0.75
        ct_DoS_Hulk = 0.5
        ct_DoS_Slowhttptest = 0.5
        ct_DoS_slowloris = 1
        ct_FTP_Patator = 0.25
        ct_Heartbleed = 0.5
        ct_Infiltration = 0.25
        ct_PortScan = 0.5
        ct_Sql_Injection = 0.75
        ct_SSH_Patator = 0.5
        ct_XSS = 0.5

        ct = [ct_benign, ct_Bot, ct_Brute_Force, ct_DDOS, ct_DoS_GoldenEye, ct_DoS_Hulk, ct_DoS_Slowhttptest,
             ct_DoS_slowloris, ct_FTP_Patator, ct_Heartbleed, ct_Infiltration, ct_PortScan, ct_Sql_Injection,
             ct_SSH_Patator, ct_XSS]

        # nt - пороговое значение аффиности для добавления в популяцию
        nt_benign = 1
        nt_Bot = 1
        nt_Brute_Force = 1
        nt_DDOS = 1
        nt_DoS_GoldenEye = 0.5
        nt_DoS_Hulk = 0.5
        nt_DoS_Slowhttptest = 1
        nt_DoS_slowloris = 0.25
        nt_FTP_Patator = 0.75
        nt_Heartbleed = 0.25
        nt_Infiltration = 0.75
        nt_PortScan = 0.75
        nt_Sql_Injection = 0.75
        nt_SSH_Patator = 0.5
        nt_XSS = 0.25

        nt = [nt_benign, nt_Bot, nt_Brute_Force, nt_DDOS, nt_DoS_GoldenEye, nt_DoS_Hulk, nt_DoS_Slowhttptest,
              nt_DoS_slowloris, nt_FTP_Patator, nt_Heartbleed, nt_Infiltration, nt_PortScan, nt_Sql_Injection,
              nt_SSH_Patator, nt_XSS]

        B_benign = 0.08
        B_Bot = 0.04
        B_Brute_Force = 0.07
        B_DDOS = 0.08
        B_DoS_GoldenEye = 0.05
        B_DoS_Hulk = 0.09
        B_DoS_Slowhttptest = 0.08
        B_DoS_slowloris = 0.03
        B_FTP_Patator = 0.1
        B_Heartbleed = 0.07
        B_Infiltration = 0.07
        B_PortScan = 0.08
        B_Sql_Injection = 0.03
        B_SSH_Patator = 0.08
        B_XSS = 0.1

        B = [B_benign, B_Bot, B_Brute_Force, B_DDOS, B_DoS_GoldenEye, B_DoS_Hulk, B_DoS_Slowhttptest,
             B_DoS_slowloris, B_FTP_Patator, B_Heartbleed, B_Infiltration, B_PortScan, B_Sql_Injection,
             B_SSH_Patator, B_XSS]

        # границы мутации
        mut_benign = [-0.01, 0.01]
        mut_Bot = [-0.00005, 0.00005]
        mut_Brute_Force = [-0.00001, 0.00001]
        mut_DDOS = [-0.00001, 0.00001]
        mut_DoS_GoldenEye = [-0.01, 0.01]
        mut_DoS_Hulk = [-0.03, 0.03]
        mut_DoS_Slowhttptest = [-0.0001, 0.0001]
        mut_DoS_slowloris = [-0.001, 0.001]
        mut_FTP_Patator = [-0.00005, 0.00005]
        mut_Heartbleed = [-0.005, 0.005]
        mut_Infiltration = [-0.00001, 0.00001]
        mut_PortScan = [-0.0005, 0.0005]
        mut_Sql_Injection = [-0.01, 0.01]
        mut_SSH_Patator = [-0.01, 0.01]
        mut_XSS = [-0.0001, 0.0001]

        mut = [mut_benign, mut_Bot, mut_Brute_Force, mut_DDOS, mut_DoS_GoldenEye, mut_DoS_Hulk, mut_DoS_Slowhttptest,
               mut_DoS_slowloris, mut_FTP_Patator, mut_Heartbleed, mut_Infiltration, mut_PortScan, mut_Sql_Injection,
               mut_SSH_Patator, mut_XSS]

        # тестовая выборка
        M = 1000

        # пороговая аффиность
        aff_benign = 0.7
        aff_Bot = 0.125
        aff_Brute_Force = 0.2
        aff_DDOS = 0.775
        aff_DoS_GoldenEye = 0.25
        aff_DoS_Hulk = 0.4
        aff_DoS_Slowhttptest = 0.4
        aff_DoS_slowloris = 0.7
        aff_FTP_Patator = 0.7
        aff_Heartbleed = 0.85
        aff_Infiltration = 0.8
        aff_PortScan = 0.8
        aff_Sql_Injection = 0.55
        aff_SSH_Patator = 0.55
        aff_XSS = 0.1

        aff = [aff_benign, aff_Bot, aff_Brute_Force, aff_DDOS, aff_DoS_GoldenEye, aff_DoS_Hulk, aff_DoS_Slowhttptest,
             aff_DoS_slowloris, aff_FTP_Patator, aff_Heartbleed, aff_Infiltration, aff_PortScan, aff_Sql_Injection,
             aff_SSH_Patator, aff_XSS]

        K, p, len_C, c_mid, te, ob, ts = main_sup(np_attack, np_no_attack, N[u - 2], ct[u - 2], nt[u - 2], 10, B[u - 2], mut[u - 2], M, aff[u - 2])

        #playsound('Filatov_Karas_GAYAZOV_BROTHER_-_Poshla_zhara_72992182.mp3')

        print(f"Тебя устраивает вероятность {p} y/n")
        #g = input()

        if p < T[u - 2]:
            continue

        else:

            c = pd.DataFrame(K)
            c.to_csv(f'Обученные модели/С множеством центров/{files[u][:-4]}/{files[u][:-4]}.csv', index=False)

            doc = docx.Document(f'Обученные модели/С множеством центров/{files[u][:-4]}/{files[u][:-4]}.docx')

            # добавляем первый параграф
            doc.add_paragraph(f'Дата {time.time()}')

            doc.add_paragraph('N = {} ct = {} nt = {} TSC = {} B = {} mut = {} M = {} '
                              'aff = {} популяция = {}'
                              .format(N[u - 2], ct[u - 2], nt[u - 2], ts, B[u - 2], mut[u - 2], M, aff[u - 2], len_C))

            doc.add_paragraph(f'вероятность {p}')
            doc.add_paragraph(f'случайная выборка {te}')
            doc.add_paragraph(f'обучающая выборка {ob}')

            doc.save(f'Обученные модели/С множеством центров/{files[u][:-4]}/{files[u][:-4]}.docx')

            u += 1



    playsound('Filatov_Karas_GAYAZOV_BROTHER_-_Poshla_zhara_72992182.mp3')
    return

### исследуем только центр
def Test1(attack, Q, aff):
    c_mid = np.sum(attack, axis=0) / len(attack)
    a = np.array([affinity(i, c_mid) for i in Q])

    A_no = np.zeros(len(a))
    f = list(a > aff)
    A_no[f] = 1
    p = np.sum(A_no)

    return p / len(A_no)
###
def Test2(attack, Q):
    ind = np.arange(len(Q))
    np.random.shuffle(ind)
    Q = Q[ind]

    #W = np.zeros(len(Q))
    W = np.zeros(1000)
    print(len(Q))

    i = 0
    while i < 100 and i < len(Q):
    #for i in range(len(Q)):
        print(i)
        a = np.array([affinity(j, Q[i]) for j in attack])

        A_no = np.zeros(len(a))
        f = list(a <= 1)
        A_no[f] = 1
        p = np.sum(A_no) / len(A_no)

        if p < 0.15:
            W[i] = 1

        i += 1

    return np.sum(W[:i]) / len(W[:i])

###
def Test3(attack, Q):
    ind = np.arange(len(Q))
    np.random.shuffle(ind)
    Q = Q[ind]

    # W = np.zeros(len(Q))
    W = np.zeros(1000)
    print(len(Q))

    i = 0
    while i < 100 and i < len(Q):
        # for i in range(len(Q)):
        print(i)
        a = np.array([affinity(j, Q[i]) for j in attack])

        a = np.sort(a)

        if a[0] > 0.05:
            W[i] = 1

        i += 1

    return np.sum(W[:i]) / len(W[:i])

### тестирование по отдельности каждой сети
def network_testing1():
    NO = np.zeros((1, 25))
    Bot = np.zeros((1, 25))
    Brute_Force = np.zeros((1, 25))
    DDoS = np.zeros((1, 25))
    DoS_GoldenEye = np.zeros((1, 25))
    DoS_Hulk = np.zeros((1, 25))
    DoS_Slowhttptest = np.zeros((1, 25))
    DoS_slowloris = np.zeros((1, 25))
    FTP_Patator = np.zeros((1, 25))
    Heartbleed = np.zeros((1, 25))
    Infiltration = np.zeros((1, 25))
    PortScan = np.zeros((1, 25))
    Sql_Injection = np.zeros((1, 25))
    SSH_Patator = np.zeros((1, 25))
    XSS = np.zeros((1, 25))

    Q = [NO, Bot, Brute_Force, DDoS, DoS_GoldenEye, DoS_Hulk, DoS_Slowhttptest,
         DoS_slowloris, FTP_Patator, Heartbleed, Infiltration, PortScan,
         Sql_Injection, SSH_Patator, XSS]

    attack = pd.DataFrame()
    no_attack = pd.DataFrame()

    directory = '/Users/User/Desktop/ДИПЛОМ/Данные по кластерам'
    files1 = os.listdir(directory)

    aff_benign = 0.7
    aff_Bot = 0.125
    aff_Brute_Force = 0.2
    #aff_DDOS = 0.25
    aff_DDOS = 0.775
    aff_DoS_GoldenEye = 0.25
    aff_DoS_Hulk = 0.4
    aff_DoS_Slowhttptest = 0.4
    aff_DoS_slowloris = 0.7
    aff_FTP_Patator = 0.45
    aff_Heartbleed = 0.85
    aff_Infiltration = 0.8
    aff_PortScan = 0.8
    aff_Sql_Injection = 0.55
    aff_SSH_Patator = 0.55
    aff_XSS = 0.1

    aff = [aff_benign, aff_Bot, aff_Brute_Force, aff_DDOS, aff_DoS_GoldenEye, aff_DoS_Hulk, aff_DoS_Slowhttptest,
           aff_DoS_slowloris, aff_FTP_Patator, aff_Heartbleed, aff_Infiltration, aff_PortScan, aff_Sql_Injection,
           aff_SSH_Patator, aff_XSS]

    for i in range(len(files1)):
        data = pd.read_csv('Данные по кластерам/' + files1[i])
        attack = pd.DataFrame()

        if i < 3:
            no_attack = pd.concat([no_attack, data])

        else:
            attack = pd.concat([attack, data])
            np_attack = attack.to_numpy()
            Q[i - 2] = np_attack

    np_no_attack = no_attack.to_numpy()
    Q[0] = np_no_attack

    C = ['С одним центром', 'С множеством центров']

    for c in range(len(C)):

        #directory = '/Users/User/Desktop/ДИПЛОМ/Обученные модели/С одним центром'
        directory = f'/Users/User/Desktop/ДИПЛОМ/Обученные модели/{C[c]}'
        files = os.listdir(directory)

        P1 = np.zeros((len(files), len(Q)))

        for i in range(len(files)):
            #data = pd.read_csv('Обученные модели/С одним центром/' + files[i] + '/' + files[i] + '.csv')
            data = pd.read_csv(f'Обученные модели/{C[c]}/' + files[i] + '/' + files[i] + '.csv')
            attack = pd.DataFrame()

            attack = pd.concat([attack, data])
            np_attack = attack.to_numpy()

            for j in range(len(Q)):
                #p1 = Test1(np_attack, Q[j], aff[i])
                #p1 = Test2(np_attack, Q[j])
                p1 = Test3(np_attack, Q[j])
                print(f"Модель с одной точкой {files[i]}, тест {files1[j + 2]}, вероятность угадывания - {p1}")
                P1[i, j] = p1

                # Test2(np_attack, Q[j])
                # Test3(np_attack, Q[j])

        time = datetime.datetime.now()

        doc = docx.Document(f'Тесты/Тестирование обученных моделей/Test3.docx')

        # добавляем первый параграф
        doc.add_paragraph(f'Дата {time}')

        table = doc.add_table(rows=(len(files) + 1), cols=(len(Q) + 2))
        # применяем стиль для таблицы
        table.style = 'Table Grid'

        for col in range(2, len(Q) + 2):
            # получаем ячейку таблицы
            cell = table.cell(0, col)
            # записываем в ячейку данные
            cell.text = str(files[col - 2])

        # заполняем таблицу данными
        for row in range(1, len(files) + 1):
            cell = table.cell(row, 0)
            cell.text = str(files[row - 1])

            cell = table.cell(row, 1)
            cell.text = str(aff[row - 1])

            for col in range(2, len(Q) + 2):
                # получаем ячейку таблицы
                cell = table.cell(row, col)
                # записываем в ячейку данные
                cell.text = str(P1[row - 1, col - 2])

        doc.save(f'Тесты/Тестирование обученных моделей/Test3.docx')

    #playsound('Filatov_Karas_GAYAZOV_BROTHER_-_Poshla_zhara_72992182.mp3')


    return


### исследуем только центр
def Test1_2(attack, Q):
    c_mid = [np.sum(i, axis=0) / len(i) for i in attack]

    F = np.zeros(15)

    for i in range(len(Q)):
        print(i)
        a = np.array([affinity(Q[i], j) for j in c_mid])
        f = np.argmin(a)
        F[f] += 1

    return F

### исследуем множество аффинностей
def Test2_2(attack, Q):
    ind = np.arange(len(Q))
    np.random.shuffle(ind)
    Q = Q[ind]
    F = np.zeros(15)

    i = 0
    while i < 100 and i < len(Q):
        print(i)
        T = np.zeros(len(attack))
        for k in range(len(attack)):
            a = np.array([affinity(j, Q[i]) for j in attack[k]])

            A_no = np.zeros(len(a))
            f = list(a <= 1)
            A_no[f] = 1
            p = np.sum(A_no) / len(A_no)

            T[k] = p

        f = np.argmax(T)
        F[f] += 1
        i += 1

    return F

### исследуем лучшую аффинность
def Test3_2(attack, Q):
    ind = np.arange(len(Q))
    np.random.shuffle(ind)
    Q = Q[ind]
    F = np.zeros(15)

    i = 0
    while i < 100 and i < len(Q):
        print(i)
        T = np.zeros(len(attack))
        for k in range(len(attack)):
            a = np.array([affinity(j, Q[i]) for j in attack[k]])

            T[k] = np.min(a)

        f = np.argmin(T)
        F[f] += 1

        i += 1

    return F


### тестирование совокупности сетей
def network_testing2():
    NO = np.zeros((1, 25))
    Bot = np.zeros((1, 25))
    Brute_Force = np.zeros((1, 25))
    DDoS = np.zeros((1, 25))
    DoS_GoldenEye = np.zeros((1, 25))
    DoS_Hulk = np.zeros((1, 25))
    DoS_Slowhttptest = np.zeros((1, 25))
    DoS_slowloris = np.zeros((1, 25))
    FTP_Patator = np.zeros((1, 25))
    Heartbleed = np.zeros((1, 25))
    Infiltration = np.zeros((1, 25))
    PortScan = np.zeros((1, 25))
    Sql_Injection = np.zeros((1, 25))
    SSH_Patator = np.zeros((1, 25))
    XSS = np.zeros((1, 25))

    Q = [NO, Bot, Brute_Force, DDoS, DoS_GoldenEye, DoS_Hulk, DoS_Slowhttptest,
         DoS_slowloris, FTP_Patator, Heartbleed, Infiltration, PortScan,
         Sql_Injection, SSH_Patator, XSS]

    attack = pd.DataFrame()
    no_attack = pd.DataFrame()

    directory = '/Users/User/Desktop/ДИПЛОМ/Данные по кластерам'
    files1 = os.listdir(directory)


    for i in range(len(files1)):
        data = pd.read_csv('Данные по кластерам/' + files1[i])
        attack = pd.DataFrame()

        if i < 3:
            no_attack = pd.concat([no_attack, data])

        else:
            attack = pd.concat([attack, data])
            np_attack = attack.to_numpy()
            Q[i - 2] = np_attack

    np_no_attack = no_attack.to_numpy()
    Q[0] = np_no_attack


    N = [NO, Bot, Brute_Force, DDoS, DoS_GoldenEye, DoS_Hulk, DoS_Slowhttptest,
         DoS_slowloris, FTP_Patator, Heartbleed, Infiltration, PortScan,
         Sql_Injection, SSH_Patator, XSS]

    no_attack = pd.DataFrame()

    C = ['С одним центром', 'С множеством центров']

    for c in range(len(C)):

        #directory = '/Users/User/Desktop/ДИПЛОМ/Обученные модели/С одним центром'
        directory = f'/Users/User/Desktop/ДИПЛОМ/Обученные модели/{C[c]}'
        files = os.listdir(directory)

        for i in range(len(files)):
            data = pd.read_csv(f'Обученные модели/{C[c]}/{files[i]}/{files[i]}.csv')
            attack = pd.DataFrame()
            attack = pd.concat([attack, data])

            np_attack = attack.to_numpy()
            N[i] = np_attack

        P1 = np.zeros((len(files), len(Q) + 2))

        for j in range(len(Q)):
            #F = Test1_2(N, Q[j])
            #F = Test2_2(N, Q[j])
            F = Test3_2(N, Q[j])

            p1 = np.sum(F[1:]) / np.sum(F)
            p2 = F[j] / np.sum(F)

            print(f"Модель с одной точкой {files[j]}, вероятность угадывания как атаки - {p1},"
                  f" вероятность отнесения к верному кластеру - {p2}")
            P1[j, 0] = p1
            P1[j, 1] = p2

            P1[j, 2:] = F

            # Test2(np_attack, Q[j])
            # Test3(np_attack, Q[j])

        time = datetime.datetime.now()

        doc = docx.Document(f'Тесты/Тестирование совокупности моделей/Test3.docx')

        # добавляем первый параграф
        doc.add_paragraph(f'Дата {time}')

        table = doc.add_table(rows=(len(files) + 1), cols=(len(Q) + 3))
        # применяем стиль для таблицы
        table.style = 'Table Grid'

        for col in range(3, len(Q) + 3):
            # получаем ячейку таблицы
            cell = table.cell(0, col)
            # записываем в ячейку данные
            cell.text = str(files[col - 3])

        # заполняем таблицу данными
        for row in range(1, len(files) + 1):
            cell = table.cell(row, 0)
            cell.text = str(files[row - 1])

            cell = table.cell(row, 1)
            cell.text = str(P1[row - 1, 0])

            cell = table.cell(row, 2)
            cell.text = str(P1[row - 1, 1])

            for col in range(3, len(Q) + 3):
                # получаем ячейку таблицы
                cell = table.cell(row, col)
                # записываем в ячейку данные
                cell.text = str(P1[row - 1, col - 1])

        doc.save(f'Тесты/Тестирование совокупности моделей/Test3.docx')

    playsound('Filatov_Karas_GAYAZOV_BROTHER_-_Poshla_zhara_72992182.mp3')
    return

### тестирование совокупности данных
def network_testing3():
    NO = np.zeros((1, 25))
    Bot = np.zeros((1, 25))
    Brute_Force = np.zeros((1, 25))
    DDoS = np.zeros((1, 25))
    DoS_GoldenEye = np.zeros((1, 25))
    DoS_Hulk = np.zeros((1, 25))
    DoS_Slowhttptest = np.zeros((1, 25))
    DoS_slowloris = np.zeros((1, 25))
    FTP_Patator = np.zeros((1, 25))
    Heartbleed = np.zeros((1, 25))
    Infiltration = np.zeros((1, 25))
    PortScan = np.zeros((1, 25))
    Sql_Injection = np.zeros((1, 25))
    SSH_Patator = np.zeros((1, 25))
    XSS = np.zeros((1, 25))

    Q = [NO, Bot, Brute_Force, DDoS, DoS_GoldenEye, DoS_Hulk, DoS_Slowhttptest,
         DoS_slowloris, FTP_Patator, Heartbleed, Infiltration, PortScan,
         Sql_Injection, SSH_Patator, XSS]

    attack = pd.DataFrame()
    no_attack = pd.DataFrame()

    directory = '/Users/User/Desktop/ДИПЛОМ/Данные по кластерам'
    files1 = os.listdir(directory)


    for i in range(len(files1)):
        data = pd.read_csv('Данные по кластерам/' + files1[i])
        attack = pd.DataFrame()

        if i < 3:
            no_attack = pd.concat([no_attack, data])

        else:
            attack = pd.concat([attack, data])
            np_attack = attack.to_numpy()
            Q[i - 2] = np_attack

    np_no_attack = no_attack.to_numpy()
    Q[0] = np_no_attack


    N = [NO, Bot, Brute_Force, DDoS, DoS_GoldenEye, DoS_Hulk, DoS_Slowhttptest,
         DoS_slowloris, FTP_Patator, Heartbleed, Infiltration, PortScan,
         Sql_Injection, SSH_Patator, XSS]

    no_attack = pd.DataFrame()



    for c in range(1, 4):

        #directory = '/Users/User/Desktop/ДИПЛОМ/Обученные модели/С одним центром'
        directory = f'/Users/User/Desktop/ДИПЛОМ/Данные взятые для модели'
        files = os.listdir(directory)

        for i in range(len(files)):
            data = pd.read_csv(f'Данные взятые для модели/{files[i]}')
            attack = pd.DataFrame()
            attack = pd.concat([attack, data])

            np_attack = attack.to_numpy()
            N[i] = np_attack

        P1 = np.zeros((len(files), len(Q) + 2))

        for j in range(len(Q)):
            if c == 1:
                F = Test1_2(N, Q[j])

            if c == 2:
                F = Test2_2(N, Q[j])
            if c == 3:
                F = Test3_2(N, Q[j])


            p1 = np.sum(F[1:]) / np.sum(F)
            p2 = F[j] / np.sum(F)

            print(f"Модель с одной точкой {files[j]}, вероятность угадывания как атаки - {p1},"
                  f" вероятность отнесения к верному кластеру - {p2}")
            P1[j, 0] = p1
            P1[j, 1] = p2

            P1[j, 2:] = F

            # Test2(np_attack, Q[j])
            # Test3(np_attack, Q[j])

        time = datetime.datetime.now()

        doc = docx.Document(f'Тесты/Тестирование совокупности данных/Test{c}.docx')

        # добавляем первый параграф
        doc.add_paragraph(f'Дата {time}')

        table = doc.add_table(rows=(len(files) + 1), cols=(len(Q) + 3))
        # применяем стиль для таблицы
        table.style = 'Table Grid'

        for col in range(3, len(Q) + 3):
            # получаем ячейку таблицы
            cell = table.cell(0, col)
            # записываем в ячейку данные
            cell.text = str(files[col - 3])

        # заполняем таблицу данными
        for row in range(1, len(files) + 1):
            cell = table.cell(row, 0)
            cell.text = str(files[row - 1])

            cell = table.cell(row, 1)
            cell.text = str(P1[row - 1, 0])

            cell = table.cell(row, 2)
            cell.text = str(P1[row - 1, 1])

            for col in range(3, len(Q) + 3):
                # получаем ячейку таблицы
                cell = table.cell(row, col)
                # записываем в ячейку данные
                cell.text = str(P1[row - 1, col - 1])

        doc.save(f'Тесты/Тестирование совокупности данных/Test{c}.docx')

    #playsound('Filatov_Karas_GAYAZOV_BROTHER_-_Poshla_zhara_72992182.mp3')
    return


if __name__ == "__main__":
    #check()

    #model_training()

    #main()


    network_testing3()

    #D_E()

    """N = [[1, 2, 1, 2],
         [2, 2, 2, 2]]

    sample = np.mean(N, axis=0)
    print(sample)"""


    #check_sample()

    #D()
    #analysis()

    #sample_cleaning()
    #separation_of_attacks()
    #data_normalization()\
    """mu = 5
    sigma = 2

    R1 = [random.normalvariate(mu, sigma) for i in range(500)]
    R2 = [random.normalvariate(mu, sigma) for i in range(500)]

    R11 = []
    R22 = []

    for i in range(len(R1)):
        if math.sqrt((R1[i] - 5) ** 2 + (R2[i] - 5) ** 2) > 2:
            R11.append(R1[i])
            R22.append(R2[i])


    mu2 = 2
    sigma = 0.15

    R3 = [random.normalvariate(mu2, sigma) for i in range(500)]
    R4 = [random.normalvariate(mu2 + 2, sigma) for i in range(500)]

    mu1 = 5.5
    mu2 = 4
    sigma = 0.3

    R5 = [random.normalvariate(mu1, sigma) for i in range(100)]
    R6 = [random.normalvariate(mu2, sigma) for i in range(100)]

    mu1 = 4.5
    mu2 = 7
    sigma = 0.2

    R7 = [random.normalvariate(mu1, sigma) for i in range(2)]
    R8 = [random.normalvariate(mu2, sigma) for i in range(2)]

    figure, axes = plt.subplots()

    plt.scatter(R3, R4, c='yellow')
    plt.scatter(R11, R22)
    plt.scatter(5, 5, marker='*', s=100, c='red')


    plt.xlim([0, 10])
    plt.ylim([0, 10])"""



    """plt.scatter(R5, R6, c='green')
    plt.scatter(R7, R8, c='black')
    
    plt.scatter(5.05, 4.9, marker='*', s=100, c='violet')"""


    plt.show()

